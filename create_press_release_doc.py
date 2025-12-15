#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_bullet_list(doc, items):
    """Add a bulleted list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title - Left-aligned (standard press release format)
title = doc.add_heading('Peter Vesterbacka & Kustaa Valtonen Join Usable as Strategic Advisors and Investors', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_format = title.runs[0].font
title_format.size = Pt(18)
title_format.bold = True
title.paragraph_format.space_after = Pt(6)

# FOR IMMEDIATE RELEASE - Left-aligned (standard format)
release_badge = doc.add_paragraph('FOR IMMEDIATE RELEASE')
release_badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
release_badge_format = release_badge.runs[0].font
release_badge_format.bold = True
release_badge_format.size = Pt(10)
release_badge.paragraph_format.space_before = Pt(3)
release_badge.paragraph_format.space_after = Pt(12)

# Dateline - Standard format: LOCATION, State – DATE (left-aligned, bold)
dateline = doc.add_paragraph()
dateline.add_run('TÓRSHAVN, Faroe Islands – December 15, 2025').bold = True
dateline.paragraph_format.space_after = Pt(12)

# First paragraph - Standard format, no italics, starts immediately after dateline
intro = doc.add_paragraph()
intro.add_run('Usable.dev').bold = True
intro.add_run(', formerly known as Flowcore, today announced that ')
intro.add_run('Peter Vesterbacka').bold = True
intro.add_run(' and ')
intro.add_run('Kustaa Valtonen').bold = True
intro.add_run(' from ')
intro.add_run('Random Ventures').bold = True
intro.add_run(' have joined the company as strategic advisors and investors. The move marks an important milestone for Usable as it scales globally following strong momentum in 2024, including being selected as a ')
intro.add_run('Top 20 company in the Slush pitching competition.').bold = True
intro.paragraph_format.space_after = Pt(12)

p2 = doc.add_paragraph('Usable is building the knowledge management layer for AI agents — enabling agents to access, share, and retain long-term organizational knowledge. The company evolved from Flowcore, where the founding team built robust event-driven data infrastructure before narrowing its focus to AI-native knowledge systems under the Usable brand.')
p2.paragraph_format.space_after = Pt(12)

# Quote - prominently displayed
quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.5)
quote.paragraph_format.right_indent = Inches(0.5)
quote.paragraph_format.space_before = Pt(6)
quote.paragraph_format.space_after = Pt(6)
quote_run = quote.add_run('"The AI agent ecosystem is moving incredibly fast, but almost everyone runs into the same limitation: agents don\'t have durable memory," said ')
quote_run.bold = True
quote_run = quote.add_run('Ólavur Ellefsen, CEO of Usable')
quote_run.bold = True
quote_run = quote.add_run('. "Usable solves that problem at the infrastructure level. With Peter and Kustaa joining us, we gain strategic insight and global experience that will help us scale much faster."')
quote_run.bold = True
quote.paragraph_format.space_after = Pt(12)

# Strategic Advisors and Investors
doc.add_heading('Strategic Advisors and Investors', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

p3 = doc.add_paragraph()
p3.add_run('Peter Vesterbacka').bold = True
p3.add_run(' is best known for his work with Rovio the Angry Birds company where he was the Mighty Eagle (CMO). He is also well known as the founder of ')
p3.add_run('Slush').bold = True
p3.add_run(', the world\'s leading startup event. He brings extensive experience in building global tech ecosystems, scaling startups, and connecting founders with international markets.')
p3.paragraph_format.space_after = Pt(6)

p4 = doc.add_paragraph()
p4.add_run('Kustaa Valtonen').bold = True
p4.add_run(' is a seasoned early-stage investor and technology advisor with deep experience in enterprise software and B2B SaaS. Through Random Ventures, he works closely with founders to scale products and organizations internationally.')
p4.paragraph_format.space_after = Pt(6)

p5 = doc.add_paragraph()
p5.add_run('Random Ventures').bold = True
p5.add_run(' is a Finest Bay Area investment company that invests globally in early-stage startups. The firm supports ambitious founders with both capital and hands-on strategic guidance across markets.')
p5.paragraph_format.space_after = Pt(12)

# Building from the Faroe Islands
doc.add_heading('Building from the Faroe Islands', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

p6 = doc.add_paragraph()
p6.add_run('Usable is headquartered in Tórshavn, Faroe Islands, a location increasingly recognized for its ')
p6.add_run('tight-knit, highly technical startup ecosystem.').bold = True
p6.add_run(' Faroese startups benefit from close collaboration, fast decision-making, and strong international orientation — qualities that have shaped Usable\'s developer-first approach from day one.')
p6.paragraph_format.space_after = Pt(12)

# What Usable Does
doc.add_heading('What Usable Does', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

p7 = doc.add_paragraph('As organizations deploy AI agents, teams are forced to choose between agents with no memory, legacy knowledge tools not built for AI, or expensive custom infrastructure.')
p7.paragraph_format.space_after = Pt(6)

p8 = doc.add_paragraph('Usable removes that complexity by providing a dedicated knowledge base for AI agents. Teams can upload or generate documentation, connect existing systems via API, structure knowledge as fragments and relationships, and expose it to any AI agent through MCP or REST APIs. The result is shared, durable memory across agents — without months of custom engineering.')
p8.paragraph_format.space_after = Pt(12)

# What's Next
doc.add_heading('What\'s Next', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

p9 = doc.add_paragraph('Usable is expanding its customer base across Europe and beyond and is actively working with developer teams building AI agents that require shared, long-term memory.')
p9.paragraph_format.space_after = Pt(12)

# Learn more
doc.add_heading('Learn more:', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

add_bullet_list(doc, [
    'Website: https://usable.dev',
    'Founder\'s blog: https://usable.dev/blog/why-were-embracing-the-usable-brand',
    'Press assets: https://usable.dev/media-kit'
])

doc.add_paragraph('')  # Add spacing
doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

# About Usable Sp/f
doc.add_heading('About Usable Sp/f', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

p10 = doc.add_paragraph('Usable (formerly Flowcore) is a Faroese technology company building developer-first knowledge management infrastructure for AI agents. The company provides the shared memory layer that enables agents to access structured, persistent knowledge across systems and sessions. Usable serves startups and developer teams globally from its base in the Faroe Islands.')
p10.paragraph_format.space_after = Pt(6)

p11 = doc.add_paragraph('Usable was founded by Ólavur Ellefsen, Julius á Rógvi Biskopstø, and Brian Bischoff. The company employs 8 full-time employees.')
p11.paragraph_format.space_after = Pt(12)

# Media Contact
doc.add_heading('Media Contact', 1)
heading = doc.paragraphs[-1]
heading.paragraph_format.space_after = Pt(6)

contact = doc.add_paragraph()
contact.paragraph_format.left_indent = Inches(0.3)
contact.paragraph_format.right_indent = Inches(0.3)
contact.paragraph_format.space_before = Pt(6)
contact.paragraph_format.space_after = Pt(6)

contact.add_run('Ólavur Ellefsen').bold = True
contact.add_run('\nChief Executive Officer\nUsable Sp/f\n')
contact.add_run('Email: ').bold = True
contact.add_run('olavur@usable.dev\n')
contact.add_run('Phone: ').bold = True
contact.add_run('+298 556600\n')
contact.add_run('Timezone: ').bold = True
contact.add_run('UTC')
contact.paragraph_format.space_after = Pt(12)

# Footer note
footer = doc.add_paragraph('---')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_format = footer.runs[0].font
footer_format.size = Pt(8)
footer.paragraph_format.space_after = Pt(6)

footer_note = doc.add_paragraph('This press release is available online at: https://usable.dev/news/peter-vesterbacka-kustaa-valtonen-join-usable-strategic-advisors-investors')
footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_note_format = footer_note.runs[0].font
footer_note_format.size = Pt(9)
footer_note_format.italic = True

doc.save('press-release-draft.docx')
print('Press release Word document created: press-release-draft.docx')
